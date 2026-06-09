"""
Génère les 3 documents fournisseur RATP pour la démo UC28.
Sortie : uc28-inspection/frontend/public/documents/
"""
import os
from pathlib import Path

OUT = Path(__file__).parent.parent / "uc28-inspection/frontend/public/documents"
OUT.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────
# 1. CR_Audit_nov_2024.docx  (python-docx)
# ─────────────────────────────────────────────────────────────
def gen_cr_audit():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
        return p

    def para(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def add_table_row(table, cells, bold_first=False):
        row = table.add_row()
        for i, (cell, text) in enumerate(zip(row.cells, cells)):
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if bold_first and i == 0:
                        run.bold = True

    # ── En-tête ──────────────────────────────────────────────
    h = doc.add_heading("BUREAU VERITAS", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x4A, 0x9F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rapport de mission d'inspection — Confidentiel")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()

    # ── Titre principal ───────────────────────────────────────
    h1 = doc.add_heading("COMPTE-RENDU D'INSPECTION", 1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h1.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Tableau récapitulatif mission ─────────────────────────
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Champ"
    hdr[1].text = "Valeur"
    for cell in hdr:
        for _p in cell.paragraphs:
            for run in _p.runs:
                run.bold = True
                run.font.size = Pt(9)

    rows = [
        ("Client", "RATP — Direction de l'Ingénierie"),
        ("Site audité", "Atelier de Sucy-en-Brie (SUC)"),
        ("Référentiel", "ISO 9001:2015"),
        ("Date de visite", "14 novembre 2024"),
        ("Auditeur", "Marc Lefèvre — Bureau Veritas"),
        ("Contact site", "Mei Lin Zhang — Responsable Qualité"),
        ("Durée", "1 journée (7 h)"),
        ("Statut rapport", "DÉFINITIF"),
    ]
    for r in rows:
        add_table_row(t, r, bold_first=True)

    doc.add_paragraph()

    # ── Résumé exécutif ───────────────────────────────────────
    heading("1. Résumé exécutif", level=2, color=(0x00, 0x4A, 0x9F))
    para(
        "L'inspection réalisée le 14 novembre 2024 à l'atelier de Sucy-en-Brie a mis en évidence "
        "1 non-conformité majeure et 1 non-conformité mineure au regard des exigences de l'ISO 9001:2015. "
        "Ces constats portent respectivement sur la maîtrise des équipements de mesure (§7.1.5) et sur "
        "la gestion des sorties non conformes (§8.7). Un plan d'actions correctives est attendu sous 30 jours.",
        size=10,
    )

    doc.add_paragraph()

    # ── Tableau des non-conformités ───────────────────────────
    heading("2. Constats d'audit", level=2, color=(0x00, 0x4A, 0x9F))

    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    headers = ["Réf.", "Clause ISO", "Niveau", "Constat", "Action requise"]
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
        for _p in t2.rows[0].cells[i].paragraphs:
            for run in _p.runs:
                run.bold = True
                run.font.size = Pt(9)

    constats = [
        (
            "NC-01",
            "§7.1.5 — Étalonnage",
            "MAJEURE",
            "3 clés dynamométriques (postes 8, 12, 15) présentent un certificat d'étalonnage échu "
            "depuis plus de 14 mois. Aucune action de remplacement ni de mise en quarantaine n'a été "
            "engagée. Le responsable métrologie est vacant depuis septembre 2024.",
            "Suspendre l'utilisation des équipements non étalonnés. Planifier étalonnage COFRAC "
            "sous 15 jours. Désigner un responsable métrologie par intérim.",
        ),
        (
            "NC-02",
            "§8.7 — Zone quarantaine",
            "MINEURE",
            "La zone de quarantaine des pièces non conformes est partiellement délimitée. "
            "Des pièces en attente de décision coexistent avec des pièces conformes dans "
            "la zone de stockage principal.",
            "Compléter le marquage au sol de la zone quarantaine. Apposer signalétique "
            "réglementaire. Vérifier registre de suivi des pièces NC.",
        ),
    ]
    for c in constats:
        add_table_row(t2, c)

    doc.add_paragraph()

    # ── Points conformes ──────────────────────────────────────
    heading("3. Points de conformité notables", level=2, color=(0x00, 0x4A, 0x9F))
    points = [
        "§7.2 — Compétences : dossiers de qualification à jour pour 18/20 techniciens.",
        "§7.5 — Maîtrise documentaire : système de gestion documentaire conforme, révisions tracées.",
        "§9.1 — Surveillance des performances : indicateurs qualité renseignés et revus mensuellement.",
    ]
    for pt in points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(pt)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Plan d'actions ────────────────────────────────────────
    heading("4. Plan d'actions correctives attendu", level=2, color=(0x00, 0x4A, 0x9F))

    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    for i, h in enumerate(["Réf. NC", "Action", "Responsable", "Échéance"]):
        t3.rows[0].cells[i].text = h
        for _p in t3.rows[0].cells[i].paragraphs:
            for run in _p.runs:
                run.bold = True
                run.font.size = Pt(9)

    actions = [
        ("NC-01", "Étalonnage COFRAC postes 8, 12, 15", "Mei Lin Zhang", "29/11/2024"),
        ("NC-01", "Nomination responsable métrologie intérim", "Direction technique", "21/11/2024"),
        ("NC-02", "Marquage sol zone quarantaine complet", "Service maintenance", "28/11/2024"),
        ("NC-02", "Mise à jour registre pièces NC", "Mei Lin Zhang", "21/11/2024"),
    ]
    for a in actions:
        add_table_row(t3, a)

    doc.add_paragraph()

    # ── Signatures ────────────────────────────────────────────
    heading("5. Signatures", level=2, color=(0x00, 0x4A, 0x9F))

    t4 = doc.add_table(rows=3, cols=2)
    t4.style = "Table Grid"
    t4.rows[0].cells[0].text = "Auditeur Bureau Veritas"
    t4.rows[0].cells[1].text = "Représentant site RATP"
    t4.rows[1].cells[0].text = "Marc Lefèvre"
    t4.rows[1].cells[1].text = "Mei Lin Zhang"
    t4.rows[2].cells[0].text = "Signature : ___________________"
    t4.rows[2].cells[1].text = "Signature : ___________________"
    for row in t4.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Bureau Veritas — 8 cours du Triangle de l'Arche, 92800 Puteaux — www.bureauveritas.fr\n"
        "Document confidentiel — Ne pas diffuser sans autorisation"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out = OUT / "CR_Audit_nov_2024.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")


# ─────────────────────────────────────────────────────────────
# 2. Procedures_etalonnage_v3.pdf  (reportlab)
# ─────────────────────────────────────────────────────────────
def gen_procedures_etalonnage():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    out = str(OUT / "Procedures_etalonnage_v3.pdf")
    doc = SimpleDocTemplate(
        out,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    BLUE = colors.HexColor("#004A9F")
    GRAY = colors.HexColor("#555555")
    LIGHT_GRAY = colors.HexColor("#F0F0F0")
    RED = colors.HexColor("#C0392B")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], textColor=BLUE, fontSize=16, spaceAfter=4)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=BLUE, fontSize=12, spaceBefore=14, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=GRAY, fontSize=10, spaceBefore=10, spaceAfter=3)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, spaceAfter=4, leading=13)
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8, textColor=GRAY)
    center = ParagraphStyle("center", parent=styles["Normal"], fontSize=9, alignment=TA_CENTER)
    bold_body = ParagraphStyle("bold_body", parent=body, fontName="Helvetica-Bold")

    story = []

    # En-tête
    story.append(Paragraph("RATP — Direction de l'Ingénierie", ParagraphStyle("org", parent=styles["Normal"], fontSize=10, textColor=GRAY, alignment=TA_RIGHT)))
    story.append(Paragraph("Atelier de Sucy-en-Brie (SUC)", ParagraphStyle("org2", parent=styles["Normal"], fontSize=9, textColor=GRAY, alignment=TA_RIGHT)))
    story.append(Spacer(1, 0.3 * cm))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("PROCÉDURE DE MAÎTRISE DES ÉQUIPEMENTS DE MESURE", title_style))
    story.append(Paragraph("Réf. : RATP-SUC-PROC-007 — Version 3 — Date : 15/01/2025", small))
    story.append(Spacer(1, 0.5 * cm))

    # Tableau de révisions
    story.append(Paragraph("Historique des révisions", h2))
    rev_data = [
        ["Version", "Date", "Objet de la révision", "Auteur"],
        ["v1", "12/03/2022", "Création initiale", "Eve Dupont"],
        ["v2", "07/09/2023", "Ajout équipements dynamométriques", "Eve Dupont"],
        ["v3", "15/01/2025", "Mise à jour fréquences + liste COFRAC", "Eve Dupont"],
    ]
    rev_table = Table(rev_data, colWidths=[1.5 * cm, 2.5 * cm, 9 * cm, 3 * cm])
    rev_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(rev_table)
    story.append(Spacer(1, 0.4 * cm))

    # 1. Objet
    story.append(Paragraph("1. Objet et domaine d'application", h1))
    story.append(Paragraph(
        "La présente procédure définit les modalités de maîtrise des équipements de mesure et de surveillance "
        "utilisés à l'atelier de Sucy-en-Brie, conformément à l'exigence <b>§7.1.5 de l'ISO 9001:2015</b>. "
        "Elle s'applique à tous les équipements dont les mesures influencent la conformité des produits et services.",
        body,
    ))

    # 2. Définitions
    story.append(Paragraph("2. Définitions", h1))
    defs = [
        ("Étalonnage", "Ensemble d'opérations établissant la relation entre les valeurs indiquées par un instrument de mesure et les valeurs correspondantes de la grandeur mesurée."),
        ("COFRAC", "Comité Français d'Accréditation — organisme national d'accréditation reconnu par l'État."),
        ("Équipement périmé", "Tout équipement dont la date limite d'étalonnage est dépassée. Ne peut être utilisé jusqu'à ré-étalonnage."),
    ]
    def_data = [[bold_body if False else k, v] for k, v in defs]
    def_data = [[k, v] for k, v in defs]
    def_table = Table(def_data, colWidths=[4 * cm, 12 * cm])
    def_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(def_table)
    story.append(Spacer(1, 0.3 * cm))

    # 3. Équipements
    story.append(Paragraph("3. Inventaire des équipements soumis à étalonnage", h1))
    story.append(Paragraph(
        "Le tableau suivant liste les équipements de mesure critiques de l'atelier SUC. "
        "Les lignes en rouge signalent des équipements dont le certificat d'étalonnage est <b>périmé</b> "
        "à la date d'édition du présent document.",
        body,
    ))
    story.append(Spacer(1, 0.2 * cm))

    equip_data = [
        ["N° Poste", "Désignation", "Marque / Modèle", "Plage", "Fréquence", "Dernier étalonnage", "Prochain", "Statut"],
        ["P-001", "Torsiomètre banc A", "Facom / E.306", "0–300 N·m", "12 mois", "02/03/2024", "02/03/2025", "OK"],
        ["P-004", "Manomètre pression", "Sauermann / SI-181", "0–10 bar", "12 mois", "14/06/2024", "14/06/2025", "OK"],
        ["P-008", "Clé dynamométrique", "Snap-on / QJDQ2N200", "40–200 N·m", "12 mois", "10/10/2023", "10/10/2024", "PÉRIMÉ"],
        ["P-010", "Multimètre électrique", "Fluke / 117", "—", "24 mois", "03/01/2024", "03/01/2026", "OK"],
        ["P-012", "Clé dynamométrique", "Facom / S.205A", "20–100 N·m", "12 mois", "05/10/2023", "05/10/2024", "PÉRIMÉ"],
        ["P-015", "Clé dynamométrique", "Facom / S.306", "60–300 N·m", "12 mois", "11/10/2023", "11/10/2024", "PÉRIMÉ"],
        ["P-017", "Pied à coulisse digital", "Mitutoyo / 500-196", "0–150 mm", "24 mois", "22/04/2024", "22/04/2026", "OK"],
        ["P-020", "Thermomètre IR", "Testo / 830-T1", "−30/+400 °C", "12 mois", "18/07/2024", "18/07/2025", "OK"],
    ]

    col_w = [1.4 * cm, 3.5 * cm, 3.2 * cm, 1.8 * cm, 1.8 * cm, 2.5 * cm, 2 * cm, 1.5 * cm]
    equip_table = Table(equip_data, colWidths=col_w)
    style_list = [
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("ALIGN", (7, 0), (7, -1), "CENTER"),
    ]
    # Lignes périmées en rouge
    for i, row in enumerate(equip_data[1:], start=1):
        if row[-1] == "PÉRIMÉ":
            style_list.append(("TEXTCOLOR", (7, i), (7, i), RED))
            style_list.append(("FONTNAME", (7, i), (7, i), "Helvetica-Bold"))
            style_list.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#FFF0F0")))
    equip_table.setStyle(TableStyle(style_list))
    story.append(equip_table)
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        "<b>Situation au 15/01/2025 :</b> 3 clés dynamométriques (P-008, P-012, P-015) présentent un certificat "
        "d'étalonnage échu depuis plus de 14 mois. Ces équipements sont à mettre en quarantaine immédiate.",
        ParagraphStyle("warn", parent=body, textColor=RED),
    ))

    # 4. Responsabilités
    story.append(Paragraph("4. Responsabilités", h1))
    story.append(Paragraph(
        "<b>Responsable Métrologie</b> (poste vacant depuis septembre 2024) : planification des étalonnages, "
        "tenue du registre, validation des certificats COFRAC, mise en quarantaine des équipements hors délai. "
        "Par intérim : Responsable Qualité (Mei Lin Zhang).",
        body,
    ))
    story.append(Paragraph(
        "<b>Techniciens de maintenance</b> : signalement immédiat de tout équipement dont la date limite approche "
        "(J−30). Ne pas utiliser un équipement dont le statut est PÉRIMÉ.",
        body,
    ))

    # 5. Références
    story.append(Paragraph("5. Références normatives", h1))
    refs = [
        "ISO 9001:2015 §7.1.5 — Ressources pour la surveillance et la mesure",
        "ISO/IEC 17025:2017 — Exigences générales pour la compétence des laboratoires",
        "NF EN ISO 6789-1:2017 — Outils à serrage par couple — Exigences de conception et de conformité",
        "COFRAC LAB REF-02 — Politique d'accréditation des laboratoires d'étalonnage",
    ]
    for r in refs:
        story.append(Paragraph(f"• {r}", body))

    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=GRAY))
    story.append(Paragraph(
        "RATP — Procédure interne confidentielle — Ne pas diffuser hors périmètre RATP / Bureau Veritas",
        ParagraphStyle("footer", parent=styles["Normal"], fontSize=7, textColor=GRAY, alignment=TA_CENTER),
    ))

    doc.build(story)
    print(f"  ✓ Procedures_etalonnage_v3.pdf")


# ─────────────────────────────────────────────────────────────
# 3. Plan_qualite_2025.pdf  (reportlab)
# ─────────────────────────────────────────────────────────────
def gen_plan_qualite():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    out = str(OUT / "Plan_qualite_2025.pdf")
    doc = SimpleDocTemplate(
        out,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    BLUE = colors.HexColor("#004A9F")
    GREEN = colors.HexColor("#1A7A4A")
    ORANGE = colors.HexColor("#D35400")
    RED = colors.HexColor("#C0392B")
    GRAY = colors.HexColor("#555555")
    LIGHT_GRAY = colors.HexColor("#F0F0F0")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], textColor=BLUE, fontSize=16, spaceAfter=4)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=BLUE, fontSize=12, spaceBefore=14, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=GRAY, fontSize=10, spaceBefore=10, spaceAfter=3)
    body = ParagraphStyle("body", parent=styles["Normal"], fontSize=9, spaceAfter=4, leading=13)
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8, textColor=GRAY)

    story = []

    # En-tête
    story.append(Paragraph("RATP — Direction de l'Ingénierie / Atelier de Sucy-en-Brie (SUC)", ParagraphStyle("org", parent=styles["Normal"], fontSize=9, textColor=GRAY, alignment=TA_RIGHT)))
    story.append(Spacer(1, 0.3 * cm))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE))
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("PLAN QUALITE 2025", title_style))
    story.append(Paragraph("Atelier de Sucy-en-Brie — Réf. RATP-SUC-PQ-2025 — Révision 0 — 03/01/2025", small))
    story.append(Paragraph("Établi par : Mei Lin Zhang (Responsable Qualité) — Validé par : Direction technique RATP", small))
    story.append(Spacer(1, 0.5 * cm))

    # 1. Contexte
    story.append(Paragraph("1. Contexte et périmètre", h1))
    story.append(Paragraph(
        "Le présent plan qualité définit les objectifs et actions prioritaires de l'atelier de Sucy-en-Brie "
        "pour l'année 2025, en réponse aux exigences de l'<b>ISO 9001:2015</b> et aux constats émis lors de "
        "l'audit Bureau Veritas du 14 novembre 2024. Il couvre l'ensemble des activités de maintenance "
        "des matériels roulants (RER A, RER B, lignes de métro).",
        body,
    ))

    # 2. Objectifs qualité
    story.append(Paragraph("2. Objectifs qualité 2025", h1))
    obj_data = [
        ["Objectif", "Indicateur", "Cible 2025", "Fréquence"],
        ["Conformité équipements de mesure", "% équipements étalonnés dans délai", "100 %", "Mensuelle"],
        ["Délai traitement NC", "Délai moyen clôture NC", "< 30 jours", "Mensuelle"],
        ["Taux de conformité ISO 9001", "Score audit interne", "> 85 %", "Trimestrielle"],
        ["Formation métrologie", "% techniciens formés", "100 %", "Annuelle"],
        ["Satisfaction client interne", "NPS interne", "> 7/10", "Semestrielle"],
    ]
    obj_table = Table(obj_data, colWidths=[5 * cm, 5 * cm, 3 * cm, 3 * cm])
    obj_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(obj_table)
    story.append(Spacer(1, 0.3 * cm))

    # 3. Plan d'actions correctives
    story.append(Paragraph("3. Plan d'actions correctives — Constats audit BV novembre 2024", h1))
    story.append(Paragraph(
        "Les actions ci-dessous découlent directement du rapport d'inspection Bureau Veritas (réf. BV-2024-SUC-112). "
        "Elles font l'objet d'un suivi mensuel en revue de direction.",
        body,
    ))
    story.append(Spacer(1, 0.2 * cm))

    actions_data = [
        ["Réf.", "Clause", "Action", "Responsable", "Échéance", "Budget", "Statut"],
        ["A-01", "§7.1.5", "Étalonnage COFRAC clés P-008, P-012, P-015", "Mei Lin Zhang", "29/11/2024", "1 200 €", "CLOTURE"],
        ["A-02", "§7.1.5", "Recrutement responsable métrologie", "DRH / Direction", "31/03/2025", "—", "EN COURS"],
        ["A-03", "§7.1.5", "Mise à jour planning étalonnage 2025", "Mei Lin Zhang", "15/01/2025", "—", "CLOTURE"],
        ["A-04", "§7.1.5", "Acquisition 2 clés dynamométriques neuves", "Direction technique", "30/06/2025", "3 800 €", "NON ALLOUE"],
        ["A-05", "§8.7", "Marquage sol zone quarantaine complet", "Service maintenance", "28/11/2024", "400 €", "PARTIEL"],
        ["A-06", "§8.7", "Installation signalétique réglementaire", "Service maintenance", "15/12/2024", "200 €", "PARTIEL"],
        ["A-07", "§8.7", "Refonte registre pièces NC (dématérialisation)", "Mei Lin Zhang", "30/04/2025", "—", "EN COURS"],
    ]

    col_w2 = [1.2 * cm, 1.5 * cm, 5 * cm, 2.8 * cm, 2.2 * cm, 1.8 * cm, 2.2 * cm]
    actions_table = Table(actions_data, colWidths=col_w2)
    style_a = [
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("ALIGN", (6, 0), (6, -1), "CENTER"),
        ("ALIGN", (5, 0), (5, -1), "RIGHT"),
    ]
    # Coloration statuts
    statut_colors = {
        "CLOTURE": GREEN,
        "EN COURS": ORANGE,
        "PARTIEL": ORANGE,
        "NON ALLOUE": RED,
    }
    for i, row in enumerate(actions_data[1:], start=1):
        statut = row[-1]
        if statut in statut_colors:
            style_a.append(("TEXTCOLOR", (6, i), (6, i), statut_colors[statut]))
            style_a.append(("FONTNAME", (6, i), (6, i), "Helvetica-Bold"))
    actions_table.setStyle(TableStyle(style_a))
    story.append(actions_table)
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        "<b>Situation au 03/01/2025 :</b> L'action A-04 (acquisition équipements) n'a pas reçu d'allocation "
        "budgétaire pour 2025. Les actions A-05 et A-06 sont partiellement réalisées — le marquage au sol "
        "couvre 60 % de la surface de la zone de quarantaine.",
        ParagraphStyle("warn", parent=body, textColor=ORANGE),
    ))

    # 4. Calendrier
    story.append(Paragraph("4. Calendrier des revues et audits 2025", h1))
    cal_data = [
        ["Événement", "Fréquence", "Responsable", "Dates prévues"],
        ["Revue de direction", "Semestrielle", "Direction technique", "Juin / Décembre 2025"],
        ["Audit interne ISO 9001", "Annuelle", "Mei Lin Zhang", "Septembre 2025"],
        ["Audit Bureau Veritas", "Annuelle", "Bureau Veritas", "Novembre 2025"],
        ["Revue indicateurs qualité", "Mensuelle", "Mei Lin Zhang", "1er jeudi du mois"],
        ["Contrôle planning étalonnage", "Trimestrielle", "Resp. Métrologie", "Jan / Avr / Juil / Oct"],
    ]
    cal_table = Table(cal_data, colWidths=[5 * cm, 3 * cm, 3.5 * cm, 5 * cm])
    cal_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(cal_table)

    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=GRAY))
    story.append(Paragraph(
        "RATP — Document interne confidentiel — Plan Qualité 2025 — Atelier SUC",
        ParagraphStyle("footer", parent=styles["Normal"], fontSize=7, textColor=GRAY, alignment=TA_CENTER),
    ))

    doc.build(story)
    print(f"  ✓ Plan_qualite_2025.pdf")


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Génération des documents fournisseur RATP...")
    gen_cr_audit()
    gen_procedures_etalonnage()
    gen_plan_qualite()
    print(f"\nFichiers dans : {OUT}")
