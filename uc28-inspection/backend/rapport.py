"""
Génération du rapport d'audit ISO 9001 au format Word (.docx).
"""
import io

from docx import Document
from docx.shared import RGBColor

_CRITICITE_LABELS = {
    "majeure": "NC MAJEURE",
    "mineure": "NC MINEURE",
    "observation": "OBSERVATION",
    "conforme": "CONFORME",
}

_CRITICITE_COLORS = {
    "majeure": RGBColor(0xDC, 0x26, 0x26),
    "mineure": RGBColor(0xEA, 0x58, 0x0C),
    "observation": RGBColor(0x25, 0x63, 0xEB),
    "conforme": RGBColor(0x16, 0xA3, 0x4A),
}


def generer_rapport(data: dict) -> bytes:
    doc = Document()

    doc.add_heading("RATP — RAPPORT D'AUDIT QUALITÉ ISO 9001:2015", 0)

    table = doc.add_table(rows=0, cols=2, style="Table Grid")
    for label, value in [
        ("Site", data["site_nom"]),
        ("Localisation", data["site_localisation"]),
        ("Date", data["date"]),
        ("Auditeur", data["auditeur"]),
        ("Référentiel", data["referentiel"]),
        ("Durée", data["duree"]),
    ]:
        row = table.add_row()
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].text = value

    doc.add_paragraph()

    doc.add_heading("Synthèse exécutive", level=1)
    constats = data["constats"]
    counts = {k: sum(1 for c in constats if c["criticite"] == k)
              for k in ("majeure", "mineure", "observation", "conforme")}
    doc.add_paragraph(
        f"L'audit réalisé sur le site {data['site_nom']} a permis de relever "
        f"{len(constats)} constat(s) : "
        f"{counts['majeure']} non-conformité(s) majeure(s), "
        f"{counts['mineure']} non-conformité(s) mineure(s), "
        f"{counts['observation']} observation(s), "
        f"{counts['conforme']} point(s) conforme(s)."
    )

    doc.add_heading("Détail des constats", level=1)
    for c in constats:
        criticite = c["criticite"]
        label = _CRITICITE_LABELS.get(criticite, criticite.upper())
        color = _CRITICITE_COLORS.get(criticite)

        p = doc.add_paragraph(style="List Bullet")
        run_label = p.add_run(f"[{label}]")
        run_label.bold = True
        if color:
            run_label.font.color.rgb = color

        clause_str = f"  {c.get('clause', '')} — " if c.get("clause") else "  "
        p.add_run(f"{clause_str}{c['constat']}")

        if c.get("action"):
            doc.add_paragraph(
                f"→ Action corrective : {c['action']}", style="List Bullet 2"
            )

    doc.add_heading("Recommandation", level=1)
    doc.add_paragraph(
        "Un suivi sous 30 jours est recommandé pour vérifier la mise en œuvre "
        "des actions correctives identifiées, notamment concernant l'étalonnage "
        "des équipements de mesure."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
