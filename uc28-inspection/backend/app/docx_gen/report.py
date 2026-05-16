"""Generate the pre-report DOCX from the Agent 3 report structure."""

from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


_NC_COLORS = {
    "nc_majeure": RGBColor(0xDC, 0x26, 0x26),
    "nc_mineure": RGBColor(0xF5, 0x9E, 0x0B),
    "observation": RGBColor(0x2C, 0x7A, 0x7B),
    "conforme": RGBColor(0x10, 0xB9, 0x81),
}

_NC_LABELS = {
    "nc_majeure": "NC MAJEURE",
    "nc_mineure": "NC MINEURE",
    "observation": "OBSERVATION",
    "conforme": "CONFORME",
}


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key} : ").bold = True
    p.add_run(value)


def generate_docx(
    inspection: dict,
    report_structure: dict,
) -> bytes:
    """Return the DOCX file as bytes."""
    doc = Document()

    # --- Page de garde ---
    title = doc.add_heading("Pré-rapport d'audit qualité", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_kv(doc, "Client", inspection.get("client_name", ""))
    _add_kv(doc, "Site", inspection.get("site_name", ""))
    _add_kv(doc, "Auditeur", inspection.get("auditor_name", ""))
    _add_kv(doc, "Référentiel", inspection.get("referential", "ISO 9001:2015"))
    _add_kv(doc, "Date", datetime.now().strftime("%d/%m/%Y"))
    doc.add_page_break()

    # --- Synthèse exécutive ---
    _add_heading(doc, "Synthèse exécutive")
    doc.add_paragraph(report_structure.get("executive_summary", ""))

    # --- Tableau de conformité ---
    _add_heading(doc, "Résumé de conformité", level=2)
    summary = report_structure.get("conformity_summary", {})
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Conforme", "Observation", "NC Mineure", "NC Majeure"]):
        hdr[i].text = label

    row = table.add_row().cells
    row[0].text = str(summary.get("conforme", 0))
    row[1].text = str(summary.get("observation", 0))
    row[2].text = str(summary.get("nc_mineure", 0))
    row[3].text = str(summary.get("nc_majeure", 0))
    doc.add_paragraph()

    # --- Constats par section ---
    _add_heading(doc, "Constats détaillés")
    for section in report_structure.get("sections", []):
        _add_heading(doc, section.get("title", ""), level=2)
        for finding in section.get("findings", []):
            classification = finding.get("classification", "")
            label = _NC_LABELS.get(classification, classification.upper())
            p = doc.add_paragraph()
            run = p.add_run(f"[{label}] ")
            run.bold = True
            color = _NC_COLORS.get(classification)
            if color:
                run.font.color.rgb = color
            if finding.get("norm_reference"):
                ref_run = p.add_run(f"{finding['norm_reference']} — ")
                ref_run.font.size = Pt(9)
            p.add_run(finding.get("reformulated_text", ""))
            if finding.get("suggested_action"):
                action_p = doc.add_paragraph(style="List Bullet")
                action_p.add_run("Action : ").italic = True
                action_p.add_run(finding["suggested_action"])
        doc.add_paragraph()

    # --- Plan d'action ---
    _add_heading(doc, "Plan d'action")
    actions = report_structure.get("action_plan", [])
    if actions:
        action_table = doc.add_table(rows=1, cols=5)
        action_table.style = "Table Grid"
        headers = action_table.rows[0].cells
        for i, h in enumerate(["Priorité", "Réf. constat", "Action", "Responsable", "Délai"]):
            headers[i].text = h
        for action in actions:
            row = action_table.add_row().cells
            row[0].text = f"P{action.get('priority', '')}"
            row[1].text = action.get("finding_ref", "")
            row[2].text = action.get("action", "")
            row[3].text = action.get("responsible", "")
            row[4].text = action.get("deadline", "")
    doc.add_paragraph()

    # --- Recommandation ---
    if report_structure.get("next_audit_recommendation"):
        _add_heading(doc, "Recommandation")
        doc.add_paragraph(report_structure["next_audit_recommendation"])

    # --- Footer (pied de page) ---
    section_obj = doc.sections[0]
    footer = section_obj.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = (
        f"UC 28 — Inspection Augmentée · "
        f"{inspection.get('auditor_name', '')} · "
        f"{datetime.now().strftime('%d/%m/%Y')}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
